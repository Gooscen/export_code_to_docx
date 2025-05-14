import os
import argparse
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

def is_match_extension(filename, extensions):
    return any(filename.lower().endswith('.' + ext.lower()) for ext in extensions)

def normalize_paths(paths):
    return [os.path.abspath(path) for path in paths]

def collect_all_files(paths):
    """
    从路径列表中获取所有文件路径（文件或文件夹混合）
    """
    all_files = set()
    for path in paths:
        abs_path = os.path.abspath(path)
        if os.path.isfile(abs_path):
            all_files.add(abs_path)
        elif os.path.isdir(abs_path):
            for root, _, files in os.walk(abs_path):
                for f in files:
                    full_path = os.path.join(root, f)
                    all_files.add(full_path)
    return all_files

def should_include(file_path, include_files, exclude_files, extensions):
    # 路径是否在 include 允许的文件列表中
    if file_path not in include_files:
        return False
    # 路径是否被排除
    if any(file_path.startswith(excluded) for excluded in exclude_files):
        return False
    # 是否匹配扩展名
    return is_match_extension(file_path, extensions)

def add_code_to_docx(doc, file_path):
    try:
        doc.add_heading(file_path, level=2)
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()

        para = doc.add_paragraph()
        run = para.add_run(content)
        run.font.name = 'Courier New'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Courier New')
        run.font.size = Pt(10)
    except Exception as e:
        print(f"⚠️ 无法读取文件 {file_path}: {e}")

def main():
    parser = argparse.ArgumentParser(description="导出指定源码文件到 Word 文档")
    parser.add_argument('--include', nargs='+', required=True, help='要包含的文件或文件夹路径')
    parser.add_argument('--exclude', nargs='*', default=[], help='要排除的文件或文件夹路径')
    parser.add_argument('--ext', nargs='+', required=True, help='要包含的扩展名，例如：go html js css txt')
    parser.add_argument('--output', default='code_export.docx', help='输出的 Word 文件名（默认：code_export.docx）')
    args = parser.parse_args()

    include_files = collect_all_files(args.include)
    exclude_files = normalize_paths(args.exclude)
    extensions = args.ext

    print(f"✅ 发现 {len(include_files)} 个候选文件，正在筛选扩展名...")

    final_files = [f for f in include_files if should_include(f, include_files, exclude_files, extensions)]

    print(f"📦 最终导出文件数：{len(final_files)}")

    doc = Document()
    doc.add_heading('代码导出文档', level=1)

    for file_path in sorted(final_files):
        add_code_to_docx(doc, file_path)

    doc.save(args.output)
    print(f"✅ 导出完成：{args.output}")

if __name__ == '__main__':
    main()